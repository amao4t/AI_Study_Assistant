# app/services/document_processor.py - Improved Document Processor

import os
import uuid
import re
import logging
import requests
from werkzeug.utils import secure_filename
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
import html
import base64
import json
from io import BytesIO

from app import db
from app.models.document import Document, DocumentChunk, Question

# Set up logging
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing uploaded documents"""
    
    def __init__(self, upload_folder, allowed_extensions):
        self.upload_folder = upload_folder
        self.allowed_extensions = allowed_extensions
        
        # Create upload folder if it doesn't exist
        if not os.path.exists(self.upload_folder):
            os.makedirs(self.upload_folder)
    
    def allowed_file(self, filename):
        """Check if the file type is allowed"""
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in self.allowed_extensions
    
    def save_file(self, file, user_id):
        """Save uploaded file and create document record"""
        if not file or not self.allowed_file(file.filename):
            return None, "Invalid file or file type not allowed"
        
        try:
            # Generate a secure filename to prevent path traversal attacks
            original_filename = secure_filename(file.filename)
            file_extension = original_filename.rsplit('.', 1)[1].lower()
            unique_filename = f"{uuid.uuid4().hex}.{file_extension}"
            file_path = os.path.join(self.upload_folder, unique_filename)
            
            # Check file size before saving
            file.seek(0, os.SEEK_END)
            file_size = file.tell()
            file.seek(0)  # Reset file pointer
            
            if file_size > 10 * 1024 * 1024:  # Limit to 10MB (increased from 5MB)
                return None, "File too large. Maximum file size is 10MB."
            
            # Save the file
            file.save(file_path)
            logger.info(f"File saved to {file_path}")
            
            # Create document record
            document = Document(
                filename=unique_filename,
                original_filename=original_filename,
                file_type=file_extension,
                file_size=file_size,
                file_path=file_path,
                user_id=user_id
            )
            
            db.session.add(document)
            db.session.commit()
            logger.info(f"Document record created with ID {document.id}")
            
            # Extract and store text chunks
            success, error = self.process_document_text(document)
            if not success:
                logger.error(f"Error processing document: {error}")
                return None, error
            
            return document, None
        except Exception as e:
            logger.exception("Error in save_file")
            db.session.rollback()
            return None, str(e)
    
    def sanitize_input(self, text):
        """Sanitize input text to prevent security issues"""
        if not text:
            return ""
            
        # Remove potentially dangerous HTML tags
        text = html.escape(text)
        
        # Remove script tags and content
        text = re.sub(r'<script.*?>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove iframe tags
        text = re.sub(r'<iframe.*?>.*?</iframe>', '', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove on* attributes (onclick, onload, etc.)
        text = re.sub(r'\s+on\w+\s*=\s*["\'][^"\']*["\']', '', text, flags=re.IGNORECASE)
        
        # Remove data URIs in attributes
        text = re.sub(r'data:[^;]*;base64,([^\'"]*)', '', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text
    
    def process_document_text(self, document):
        """Extract text from document and create chunks"""
        file_path = document.file_path
        file_type = document.file_type
        
        try:
            logger.info(f"Processing document: {document.id}, type: {file_type}")
            
            # Extract text based on file type
            if file_type == 'pdf':
                try:
                    text = self._extract_text_from_pdf(file_path)
                except ImportError:
                    logger.warning("PyPDF2 not available, using basic text extraction")
                    text = self._extract_text_basic(file_path)
            elif file_type == 'docx':
                try:
                    import docx
                    text = self._extract_text_from_docx(file_path)
                except ImportError:
                    logger.warning("python-docx not available, using basic text extraction")
                    text = self._extract_text_basic(file_path)
            elif file_type == 'txt':
                text = self._extract_text_from_txt(file_path)
            else:
                return False, f"Unsupported file type: {file_type}"
            
            # Sanitize the extracted text
            text = self.sanitize_input(text)
            
            text_length = len(text)
            logger.info(f"Text extracted and sanitized, length: {text_length} characters")
            
            # Truncate very long texts to prevent memory issues
            max_text_length = 100000
            if text_length > max_text_length:
                logger.warning(f"Text too long ({text_length} chars), truncating to {max_text_length}")
                text = text[:max_text_length]
                text_length = max_text_length
            
            # Improved chunking with better overlap handling
            chunk_size = 1000
            overlap = 100
            
            # Increased maximum number of chunks
            MAX_CHUNKS = 100
            
            # Delete any existing chunks for this document
            existing_chunks = DocumentChunk.query.filter_by(document_id=document.id).all()
            for chunk in existing_chunks:
                db.session.delete(chunk)
            db.session.commit()
            
            # Create chunks
            chunks = []
            start = 0
            chunk_index = 0
            
            while start < text_length and chunk_index < MAX_CHUNKS:
                # Calculate chunk end position
                end = min(start + chunk_size, text_length)
                
                # Try to find a good break point (sentence end)
                if end < text_length:
                    # Look for sentence endings (.!?) with possible trailing spaces
                    match = re.search(r'[.!?]\s*', text[max(end-50, start):end])
                    if match:
                        end = max(end-50, start) + match.end()
                
                # Extract chunk
                chunk_text = text[start:end].strip()
                
                # Skip empty chunks
                if not chunk_text:
                    start = end
                    continue
                
                # Add chunk to list
                chunks.append({
                    'chunk_text': chunk_text,
                    'chunk_index': chunk_index,
                    'document_id': document.id
                })
                
                # Move to next chunk with proper overlap
                start = end - overlap
                chunk_index += 1
            
            if chunk_index >= MAX_CHUNKS and start < text_length:
                logger.warning(f"Reached maximum chunk limit ({MAX_CHUNKS}), truncating document")
            
            # Process chunks in parallel
            return self._process_chunks_parallel(document.id, chunks)
        
        except Exception as e:
            logger.exception("Error in process_document_text")
            db.session.rollback()
            return False, str(e)
    
    def _process_chunks_parallel(self, document_id, chunks):
        """Process document chunks in parallel for better performance"""
        if not chunks:
            return True, None
            
        try:
            logger.info(f"Processing {len(chunks)} chunks in parallel for document {document_id}")
            
            # Calculate optimal number of workers based on available CPUs
            max_workers = min(8, os.cpu_count() or 4)
            logger.info(f"Using {max_workers} workers for parallel processing")
            
            # Process chunks in parallel batches
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Submit tasks
                futures = []
                for chunk_data in chunks:
                    futures.append(executor.submit(self._create_chunk, chunk_data))
                
                # Process results
                for future in as_completed(futures):
                    try:
                        success = future.result()
                        if not success:
                            logger.warning(f"Failed to create a chunk for document {document_id}")
                    except Exception as e:
                        logger.exception(f"Error in chunk processing: {e}")
            
            logger.info(f"Parallel chunk processing complete for document {document_id}")
            return True, None
            
        except Exception as e:
            logger.exception(f"Error in parallel chunk processing: {e}")
            db.session.rollback()
            return False, str(e)
    
    def _create_chunk(self, chunk_data):
        """Create a single document chunk (for parallel processing)"""
        try:
            # Create new session for this thread to avoid conflicts
            from app import create_app
            app = create_app()
            with app.app_context():
                from app import db
                from app.models.document import DocumentChunk
                
                # Create and save chunk with thread-local session
                doc_chunk = DocumentChunk(
                    chunk_text=chunk_data['chunk_text'],
                    chunk_index=chunk_data['chunk_index'],
                    document_id=chunk_data['document_id']
                )
                
                db.session.add(doc_chunk)
                db.session.commit()
                
                return True
            
        except Exception as e:
            logger.exception(f"Error creating chunk: {e}")
            # Rollback in case of error
            try:
                db.session.rollback()
            except:
                pass
            return False
    
    def _extract_text_from_pdf(self, file_path):
        """Extract text from PDF file with improved processing for larger documents"""
        import PyPDF2
        text = ""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f, strict=False)
                # Increased limit of pages to process (max 50 pages)
                pages_to_process = min(len(pdf_reader.pages), 50)
                logger.info(f"Processing {pages_to_process} pages from PDF")
                
                # Using ThreadPoolExecutor for parallel page processing
                with ThreadPoolExecutor(max_workers=4) as executor:
                    # Create tasks to process each page
                    future_to_page = {
                        executor.submit(self._extract_page_text, pdf_reader, i): i 
                        for i in range(pages_to_process)
                    }
                    
                    # Collect results
                    page_texts = [""] * pages_to_process
                    success_count = 0
                    
                    for future in as_completed(future_to_page):
                        page_num = future_to_page[future]
                        try:
                            page_text = future.result()
                            if page_text and len(page_text.strip()) > 0:
                                page_texts[page_num] = f"\n--- Page {page_num+1} ---\n{page_text}\n"
                                success_count += 1
                        except Exception as e:
                            logger.warning(f"Error extracting text from page {page_num}: {e}")
                    
                    # Check if we extracted any meaningful text
                    if success_count == 0 and pages_to_process > 0:
                        logger.warning("Failed to extract text from any pages, trying OCR fallback")
                        try:
                            # Try OCR as a fallback
                            return self.extract_text_from_pdf_with_ocr(file_path)
                        except Exception as ocr_error:
                            logger.warning(f"OCR fallback also failed: {ocr_error}")
                    
                    # Combine all text in page order
                    text = "".join(page_texts)
                    
            # If we got very little text, try OCR as a fallback
            if len(text.strip()) < 100 and pages_to_process > 0:
                logger.info("Very little text extracted, trying OCR fallback")
                try:
                    ocr_text = self.extract_text_from_pdf_with_ocr(file_path)
                    if len(ocr_text.strip()) > len(text.strip()):
                        return ocr_text
                except Exception as ocr_error:
                    logger.warning(f"OCR fallback failed: {ocr_error}")
                    
            return text
        except Exception as e:
            logger.exception(f"Error extracting text from PDF: {e}")
            # Try OCR if regular extraction fails
            try:
                logger.info("Regular PDF extraction failed, trying OCR fallback")
                return self.extract_text_from_pdf_with_ocr(file_path)
            except Exception as ocr_error:
                logger.warning(f"OCR fallback also failed: {ocr_error}")
                # Return empty string to fail gracefully
                return ""
    
    def _extract_page_text(self, pdf_reader, page_num):
        """Extract text from a single PDF page with improved error handling"""
        try:
            page = pdf_reader.pages[page_num]
            extracted_text = page.extract_text() or ""
            
            # Additional cleaning for better quality
            extracted_text = extracted_text.replace('\n\n', '\n').strip()
            
            # Check if the text contains actual content (not just whitespace or common PDF artifacts)
            if len(extracted_text.strip()) < 10:
                # Try alternate extraction method
                extracted_text = self._extract_page_text_alternate(page) or ""
                
            return extracted_text
        except Exception as e:
            logger.warning(f"Failed to extract text from page {page_num}: {e}")
            return ""
            
    def _extract_page_text_alternate(self, page):
        """Alternative method to extract text from a PDF page when the primary method fails"""
        try:
            # Try to extract text using a different approach (accessing raw stream data)
            if '/Contents' in page and hasattr(page['/Contents'], 'get_data'):
                data = page['/Contents'].get_data()
                if data:
                    # Simple text extraction from content stream
                    text = ""
                    data_str = data.decode('latin-1', errors='replace')
                    pattern = r'\((.*?)\)'
                    matches = re.findall(pattern, data_str)
                    if matches:
                        text = " ".join(matches)
                        return text
            return ""
        except Exception as e:
            logger.warning(f"Alternative text extraction failed: {e}")
            return ""
    
    def _extract_text_from_docx(self, file_path):
        """Extract text from DOCX file with improved handling"""
        import docx
        try:
            doc = docx.Document(file_path)
            
            # Increased paragraph limit
            max_paragraphs = 200  # Increased from 100 to 200
            
            text = []
            
            # Process paragraphs with limit
            paragraphs = doc.paragraphs[:max_paragraphs] if len(doc.paragraphs) > max_paragraphs else doc.paragraphs
            for paragraph in paragraphs:
                if paragraph.text:
                    text.append(paragraph.text)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            logger.exception("Error extracting text from DOCX")
            # Return empty string to fail gracefully
            return ""
    
    def _extract_text_from_txt(self, file_path):
        """Extract text from TXT file with better encoding handling"""
        try:
            # Increased character limit from 50,000 to 100,000
            max_chars = 100000
            encodings = ['utf-8', 'latin-1', 'ascii', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                        text = f.read(max_chars)
                        return text
                except UnicodeDecodeError:
                    continue
            
            # Fallback to binary mode with replacement
            with open(file_path, 'rb') as f:
                binary_data = f.read(max_chars)
                return binary_data.decode('utf-8', errors='replace')
                
        except Exception as e:
            logger.exception("Error extracting text from TXT")
            # Return empty string to fail gracefully
            return ""
    
    def _extract_text_basic(self, file_path):
        """Fallback text extraction - try to read as text with various encodings"""
        encodings = ['utf-8', 'latin-1', 'ascii', 'cp1252']
        max_chars = 100000  # Increased from 50000 to 100000
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                    return f.read(max_chars)
            except:
                continue
                
        # Fallback to binary read with replacement character
        try:
            with open(file_path, 'rb') as f:
                binary_data = f.read(max_chars)
                return binary_data.decode('utf-8', errors='replace')
        except:
            return ""
    
    def get_document_text(self, document_id):
        """Get text from document chunks"""
        document = Document.query.get(document_id)
        if not document:
            return None, "Document not found"
        
        document.update_last_accessed()
        
        try:
            # Get chunks
            chunks = DocumentChunk.query.filter_by(document_id=document_id).order_by(DocumentChunk.chunk_index).all()
            
            if not chunks:
                # Fallback to extraction if no chunks
                return self._extract_document_text(document), None
            
            # Combine chunks - increased maximum number of chunks to combine
            MAX_CHUNKS_TO_COMBINE = 100  # Increased from 50 to 100
            if len(chunks) > MAX_CHUNKS_TO_COMBINE:
                logger.warning(f"Limiting to {MAX_CHUNKS_TO_COMBINE} chunks for text retrieval (was {len(chunks)})")
                chunks = chunks[:MAX_CHUNKS_TO_COMBINE]
            
            # Combine chunks
            text = "\n".join(chunk.chunk_text for chunk in chunks)
            return text, None
            
        except Exception as e:
            logger.exception(f"Error getting document text: {e}")
            return None, str(e)
    
    def _extract_document_text(self, document):
        """Fallback to extract text directly from file"""
        try:
            file_type = document.file_type
            file_path = document.file_path
            
            if file_type == 'pdf':
                try:
                    import PyPDF2
                    return self._extract_text_from_pdf(file_path)
                except ImportError:
                    return self._extract_text_basic(file_path)
                    
            elif file_type == 'docx':
                try:
                    import docx
                    return self._extract_text_from_docx(file_path)
                except ImportError:
                    return self._extract_text_basic(file_path)
                    
            elif file_type == 'txt':
                return self._extract_text_from_txt(file_path)
                
            else:
                return "Unsupported file type"
                
        except Exception as e:
            logger.exception(f"Error in _extract_document_text: {e}")
            return "Error extracting text"
    
    def delete_document(self, document_id, user_id):
        """Delete document and its file"""
        document = Document.query.filter_by(id=document_id, user_id=user_id).first()
        if not document:
            return False, "Document not found or you don't have permission"
        
        try:
            # Delete the file
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Delete related questions first
            questions = Question.query.filter_by(document_id=document_id).all()
            for question in questions:
                db.session.delete(question)
            
            # Delete document chunks
            chunks = DocumentChunk.query.filter_by(document_id=document_id).all()
            for chunk in chunks:
                db.session.delete(chunk)
                
            # Delete the document record
            db.session.delete(document)
            db.session.commit()
            
            return True, None
        
        except Exception as e:
            db.session.rollback()
            logger.exception(f"Error deleting document: {e}")
            return False, str(e)
    
    def summarize_document(self, document_id):
        """Summarize a document using Claude API"""
        document = Document.query.get(document_id)
        if not document:
            return None, "Document not found"
        
        try:
            # Get document text
            text, error = self.get_document_text(document_id)
            if error:
                return None, f"Error getting document text: {error}"
            
            # Use Claude API to summarize
            from flask import current_app
            api_key = current_app.config['ANTHROPIC_API_KEY']
            
            if not api_key:
                return None, "API key not configured"
            
            # Maximum length for text to summarize - increased limit
            max_length = 20000  # Increased from 10000
            if len(text) > max_length:
                logger.warning(f"Text too long ({len(text)} chars) for summarization, truncating to {max_length}")
                text = text[:max_length]
            
            # API endpoint and headers
            api_url = "https://api.anthropic.com/v1/messages"
            headers = {
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json"
            }
            
            # Create prompt for document summarization
            messages = [
                {
                    "role": "user",
                    "content": f"""Please provide a comprehensive summary of the following document.
                    Organize the summary with clear headings and bullet points where appropriate.
                    Capture all the main points, key arguments, important details, and conclusions.
                    The summary should be thorough but concise, focusing on the essential information.
                    
                    DOCUMENT TO SUMMARIZE:
                    {text}
                    """
                }
            ]
            
            # Prepare API request
            payload = {
                "model": "claude-3-haiku-20240307",
                "max_tokens": 1500,
                "temperature": 0.3,
                "messages": messages
            }
            
            # Make API call
            logger.info(f"Sending summarization request to Claude API for document {document_id}")
            response = requests.post(api_url, headers=headers, json=payload)
            
            # Check for errors
            if response.status_code != 200:
                logger.error(f"Claude API error: {response.status_code} - {response.text}")
                return None, f"Error from Claude API: {response.status_code}"
            
            # Extract summary
            result = response.json()
            if 'content' in result and len(result['content']) > 0:
                summary = result['content'][0]['text'].strip()
                logger.info(f"Successfully summarized document {document_id}")
                return summary, None
            else:
                logger.error("Invalid response from Claude API")
                return None, "Failed to generate summary"
            
        except Exception as e:
            logger.exception(f"Error in summarize_document: {str(e)}")
            return None, str(e)
    
    def extract_text_from_image(self, image_path):
        """Extract text from images using OCR"""
        try:
            # Import here to avoid requiring these libraries for basic functionality
            import pytesseract
            from PIL import Image
            
            logger.info(f"Extracting text from image: {image_path}")
            img = Image.open(image_path)
            
            # Preprocess image for better OCR results (optional)
            # img = img.convert('L')  # Convert to grayscale
            
            # Extract text using Tesseract OCR
            text = pytesseract.image_to_string(img)
            
            # Clean up the text
            text = self.sanitize_input(text)
            
            logger.info(f"Successfully extracted {len(text)} characters from image")
            return text
        except Exception as e:
            logger.exception(f"Error extracting text from image: {str(e)}")
            return ""
    
    def extract_text_from_pdf_with_ocr(self, pdf_path):
        """Extract text from PDF with OCR for image-based PDFs"""
        try:
            # Import here to avoid requiring these libraries for basic functionality
            import pytesseract
            from pdf2image import convert_from_path
            
            logger.info(f"Converting PDF to images using OCR: {pdf_path}")
            # Convert PDF to images
            pages = convert_from_path(pdf_path, 300)  # DPI 300
            
            # Process each page
            text = ""
            page_count = len(pages)
            logger.info(f"Processing {page_count} pages with OCR")
            
            for i, page in enumerate(pages[:10]):  # Limit to first 10 pages for performance
                logger.info(f"OCR processing page {i+1}/{min(page_count, 10)}")
                
                # Extract text using OCR
                page_text = pytesseract.image_to_string(page)
                text += f"\n\n--- Page {i+1} ---\n\n{page_text}"
            
            # Clean the extracted text
            text = self.sanitize_input(text)
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF with OCR")
            return text
        except ImportError as e:
            logger.warning(f"OCR libraries not available: {e}")
            return ""
        except Exception as e:
            logger.exception(f"Error extracting text from PDF with OCR: {str(e)}")
            return ""
    
    def analyze_image_content(self, image_path, api_key=None):
        """Analyze image content using Claude Vision API"""
        try:
            logger.info(f"Analyzing image content: {image_path}")
            
            # If no API key provided, try to get from current_app
            if not api_key:
                from flask import current_app
                api_key = current_app.config.get('ANTHROPIC_API_KEY')
                
            if not api_key:
                logger.error("No API key available for image analysis")
                return "Error: API key not configured for image analysis"
            
            # Read and encode the image
            with open(image_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            # Prepare API call
            headers = {
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json"
            }
            
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Describe this image in detail and extract any text visible."},
                        {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64_image}}
                    ]
                }
            ]
            
            # Call Claude Vision API
            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers,
                json={"model": "claude-3-opus-20240229", "messages": messages, "max_tokens": 1000}
            )
            
            response.raise_for_status()
            result = response.json()
            
            analysis = result['content'][0]['text'] if 'content' in result and result['content'] else ""
            logger.info(f"Successfully analyzed image, generated {len(analysis)} characters of description")
            
            return analysis
            
        except Exception as e:
            logger.exception(f"Error analyzing image: {str(e)}")
            return f"Error analyzing image: {str(e)}"
    
    def process_document_with_images(self, document_id):
        """Process a document that may contain images using OCR and image analysis"""
        document = Document.query.get(document_id)
        if not document:
            return False, "Document not found"
        
        try:
            # Extract text based on file type
            if document.file_type == 'pdf':
                # Try normal extraction first
                text = self._extract_text_from_pdf(document.file_path)
                
                # If little text was extracted, likely an image-based PDF
                if len(text.strip()) < 100:
                    logger.info("PDF appears to be image-based, using OCR")
                    text = self.extract_text_from_pdf_with_ocr(document.file_path)
            else:
                # For other document types, use existing methods
                text = self._extract_document_text(document)
            
            # Sanitize and chunk the text
            text = self.sanitize_input(text)
            
            # Delete existing chunks
            existing_chunks = DocumentChunk.query.filter_by(document_id=document.id).all()
            for chunk in existing_chunks:
                db.session.delete(chunk)
            db.session.commit()
            
            # Create new chunks
            chunks = []
            chunk_size = 1000
            overlap = 100
            text_length = len(text)
            start = 0
            chunk_index = 0
            MAX_CHUNKS = 100
            
            while start < text_length and chunk_index < MAX_CHUNKS:
                end = min(start + chunk_size, text_length)
                
                # Find good break point
                if end < text_length:
                    match = re.search(r'[.!?]\s*', text[max(end-50, start):end])
                    if match:
                        end = max(end-50, start) + match.end()
                
                chunk_text = text[start:end].strip()
                
                if not chunk_text:
                    start = end
                    continue
                
                chunks.append({
                    'chunk_text': chunk_text,
                    'chunk_index': chunk_index,
                    'document_id': document.id
                })
                
                start = end - overlap
                chunk_index += 1
            
            # Process chunks in parallel
            success, error = self._process_chunks_parallel(document.id, chunks)
            if not success:
                return False, error
                
            return True, None
            
        except Exception as e:
            logger.exception(f"Error processing document with images: {str(e)}")
            db.session.rollback()
            return False, str(e)